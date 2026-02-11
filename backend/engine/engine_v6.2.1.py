#!/usr/bin/env python3
# -*- coding: utf-8 -*-

# ============================================================
#  SpectrometerGPT engine v6.2.1
#  - Externt linjebibliotek (JSON)
#  - Fluorescent-lampa-detektor
#  - Engine6-union (dataset + band)
#  - Förbättrad DOCX-rapport (v6.1.5-style)
#  - v6.2.1: buggfixar (graf-Y, fallback-docx, SHA256, auto-läge-gissning, m.m.)
# ============================================================

import os, json, math, csv, argparse, hashlib, datetime, sys
import numpy as np
import pandas as pd
from PIL import Image, ImageDraw, ImageFont
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import shutil

# ------------------------------------------------------------------
#  Tidiga stubbar (för kompatibilitet vid import)
# ------------------------------------------------------------------

def build_analysis_log_stub(*args, **kwargs):
    return ""

def build_dynamic_abstract_stub(*args, **kwargs):
    return ""

# Bind names tidigt; riktiga definitioner nedan skriver över dessa.
build_analysis_log = build_analysis_log_stub
build_dynamic_abstract = build_dynamic_abstract_stub

# ------------------------------------------------------------------
#  Tidiga riktiga definitioner (så __main__/DOCX kan kalla dem)
# ------------------------------------------------------------------

def build_dynamic_abstract(meta, matches, peak_count, lib_atoms):
    from collections import Counter
    words_target = 520
    n_lines = len(matches or [])
    try:
        species_counts = Counter([m.get("species","").split()[0] for m in (matches or [])])
    except Exception:
        species_counts = Counter()
    top_species = ", ".join([f"{s} (n={c})" for s,c in species_counts.most_common(6)]) if species_counts else "—"

    prim = sorted(matches or [], key=lambda z: (-z.get('sigma',3), -z.get('score',0)))[:min(8, len(matches or []))]
    prim_str = "; ".join([f"{chem_symbol(m.get('species',''))} {m.get('ref_nm',0.0):.1f} nm (σ={m.get('sigma',3)})" for m in prim]) if prim else "—"

    offset_txt = meta.get("Offset (median Δλ)", "0.000 nm") if isinstance(meta, dict) else "0.000 nm"
    mode_txt   = meta.get("Spektrumläge", "auto") if isinstance(meta, dict) else "auto"

    refset = {round(m.get("ref_nm",0.0),3) for m in (matches or [])}
    hg_set = {404.656, 435.833, 491.604, 546.074, 576.959, 579.066}
    ne_red_cnt = sum(1 for m in (matches or []) if str(m.get("species","")).split()[0]=="Ne" and 650.0 <= m.get("ref_nm",0.0) <= 703.5)
    ar_ok = len([x for x in refset if 696<=x<=772.5]) >= 2
    na_ok = {588.995, 589.592}.issubset(refset)
    hg_ok = len(hg_set.intersection(refset)) >= 3
    he_ok = len({447.15, 471.31, 492.19, 501.57, 706.52}.intersection(refset)) >= 2

    bands = []
    if na_ok: bands.append("Na-dublett vid 589 nm")
    if hg_ok: bands.append("Hg-grupp i blå-grön region")
    if ne_red_cnt >= 4: bands.append("Ne-kluster i rött")
    if ar_ok: bands.append("Ar-kluster 696–772 nm")
    if he_ok: bands.append("He-signatur")
    for mol in ["C2","CN","CH","O2","H2O","N2","N2+"]:
        if any(str(m.get("species","")).split()[0]==mol for m in (matches or [])):
            bands.append(mol)
    bands_str = ", ".join(bands) if bands else "—"

    parts = []
    parts.append("Detta dokument sammanfattar en automatisk spektralanalys där både bandbild och numerisk rådata används för att extrahera och identifiera karakteristiska spektrallinjer. Motorn arbetar i autoläge och kör parallella pipelines för emission och absorption; den väg som ger bäst helhetspoäng och flest robusta träffar väljs för rapporteringen. Pixel-till-nanometer-relation fås antingen direkt från datasetets nm-kolumn eller via instrumentets fabriksfit. Intensiteten normaliseras med I=max(R,G,B) och utvärderas därefter med flerstegsdetection av lokala toppar och platåer.")
    parts.append(f"I denna körning hittades totalt {n_lines} etiketter efter kvalitetskontroller och gating. Antalet kandidattoppar före matchning var {peak_count}, vilket ger en god bild av SNR och struktur i datat. De högst prioriterade identifieringarna omfattar {prim_str}. Vid en samlad artsammanställning dominerar: {top_species}. Kluster/band som bedömts relevanta i materialet: {bands_str}. Den robusta offset-uppskattningen (median av Δλ efter IQR-trim) blev {offset_txt}, och det valda analysläget var {mode_txt}.")
    parts.append("Urvalet av linjer styrs av en kombination av närhet till referens (inom våglängdstolerans), relativ toppstyrka och domänregler som gynnar fysiskt rimliga kombinationer (t.ex. att Na-dublettens båda komponenter uppträder tillsammans eller att Hg-gruppen visar flera av sina välkända linjer). För att undvika falska positiva vid breda toppar eller mättnad används ett platåmedvetet kriterium: när intensiteten når en hög nivå identifieras först en sammanhängande platå, därefter används centroid i ett lokalt fönster för att placera linjecentrum även när platån är smal. Samtidigt begränsar ett efterföljande korstest kravet på en tydlig lokal topp i närheten av den föreslagna våglängden.")
    parts.append("Kalibreringen betraktas som styv över våglängdsområdet men kan visa små systematiska avvikelser i det röda området. Därför stöds en försiktig, data-villkorad finjustering: om flera stabila röda referenser hittas beräknas en svag linjär korrektion som reducerar residualerna utan att ”vrida om” den övergripande skalan. I praktiken innebär detta att välkända röda linjer – som Na-dublett, O I 630 nm, H-serien i rött eller K-dublett kring 770 nm – lättare hamnar inom tolerans även när kurvans röda del är något förskjuten.")
    parts.append("Efter den inledande matchningen följer ett antal kontrollsteg. Gating begränsar hur många arter som får dominera inom ett 50-nm-fönster och globalt, för att undvika övermärkning. Domänreglerna beskriver enkla fysikargument som ger poängjusteringar snarare än hårda avslag, så att svaga men inte orimliga kandidater fortfarande kan redovisas, om än med låg prioritet. Slutligen kräver kvalitetskontrollen en lokal topp i närheten av varje etikett; etiketter utan sådan evidens noteras i en logg och filtreras bort.")
    parts.append("Rapportens figurer består av en annoterad bandbild och en motsvarande spektrumprofil. Bandbildens etiketter ritas med korta guider och kemiska symboler; om flera etiketter ligger tätt staplas de i tiers för att minska överlapp. Spektrumprofilen visar normaliserad intensitet med markörer vid referenserna. Den fullständiga linjetabellen redovisar mät-nm, korrigerad nm, pixelposition (om tillämpligt), källa (band/data/båda), referens-nm, residual Δλ, relativ styrka, σ-nivå och eventuella flaggor.")
    parts.append("Sammantaget ger metodiken en reproducerbar pipeline som balanserar upptäcktsförmåga och försiktighet. Den klarar både höga smala toppar och bredare band samt hanterar överexponeringar utan att låsa sig vid binära beslut. De resultat som presenteras i denna rapport bör tolkas som ett bästa-förslag givet den tillgängliga datan; för experimentella uppställningar med varierande optik, gitter och fokus rekommenderas att man kompletterar med egna kalibreringsmätningar över hela området, särskilt i rött/när-IR där små skalfel kan ge märkbara utslag i matchningen.")
    text = " ".join(parts)
    while len(text.split()) < words_target:
        parts.append("Slutligen kan resultaten även jämföras mot alternativa källor, som lampor med kända emissionslinjer (Hg, Ne, Ar, He) eller spektra från välkarakteriserade objekt. En sådan jämförelse förbättrar både tolkningen av residualer och säkerställer att algoritmens toleranser är rimligt valda för aktuell upplösning och signal-till-brusförhållande.")
        text = " ".join(parts)
    return text

def build_analysis_log(nm_axis, I_norm, peak_idx, matches, offset, qc_log, mode_label):
    import numpy as _np
    peaks = [(float(nm_axis[i]), float(I_norm[i])) for i in peak_idx]
    peaks = sorted(peaks, key=lambda t: t[1], reverse=True)[:50]
    try:
        segs = plateau_segments(I_norm, nm_axis, rel_top=0.92, rel_eps=0.012)
    except Exception:
        segs = []
    lines = []
    lines.append(f"Körning: läge={mode_label}, robust offset (median Δλ)={offset:+.3f} nm.")
    lines.append(f"Antal kandidattoppar före matchning: {len(peak_idx)}.")
    if segs:
        lines.append(f"Platåsegment (rel_top=0.92): {len(segs)} st:")
        for (nm_lo, nm_hi, nm_mid, L, R) in segs[:20]:
            lines.append(f"  • {nm_lo:.2f}–{nm_hi:.2f} nm (centrum {nm_mid:.2f} nm, bredd {nm_hi-nm_lo:.2f} nm, pix {L}-{R})")
    if peaks:
        lines.append("Topp-peakar (nm, I_norm):")
        for nmv, iv in peaks[:20]:
            lines.append(f"  • {nmv:.2f} nm — {iv:.3f}")
    lines.append(f"Etiketter efter QC/gating: {len(matches)} st:")
    for m in sorted(matches, key=lambda z: z.get('score',0.0), reverse=True)[:60]:
        lines.append(f"  • {m.get('species','')} ref {m.get('ref_nm',0.0):.2f} nm; korr {m.get('corrected_nm',m.get('nm_meas',0.0)):.2f} nm; Δλ={m.get('delta_nm',0.0):+.2f}; σ={m.get('sigma',3)}; score={m.get('score',0.0):.1f}; flags={','.join(m.get('flags',[])) if m.get('flags') else ''}")
    if qc_log:
        lines.append("QC-logg:")
        for q in qc_log:
            lines.append("  • " + str(q))
    return "\n".join(lines)

# ------------------------------------------------------------------
#  CLI-hjälpare
# ------------------------------------------------------------------

def _ALLOW_CLI():
    import sys
    return any(a.startswith("--") for a in sys.argv[1:])


# === v6.2 additions: external library loader + fluorescent detector ===

__version__ = 'v6.2.1'
DEFAULT_K=7
DEFAULT_THRESHOLD=0.06
DEFAULT_MIN_DIST_NM=2.0
DEFAULT_TOL_NM=2.0
DEFAULT_MODE="auto"

A2=8.45736375e-06
A1=0.406760986
A0=375.834988

MAX_LABELS=56
MAX_LABELS_PER_50NM=8
MAX_SPECIES_PER_50NM=4
MAX_SPECIES_GLOBAL=14

# ------------------------------------------------------------------
#  Allmänna hjälpfunktioner
# ------------------------------------------------------------------

def sha256_of_file(path):
    h=hashlib.sha256()
    with open(path,"rb") as f:
        for ch in iter(lambda:f.read(1<<20), b""):
            h.update(ch)
    return h.hexdigest()

# v6.2.1 FIX: säker variant som klarar None / saknad fil
def sha256_safe(path):
    try:
        if path and os.path.exists(path):
            return sha256_of_file(path)[:16] + "…"
    except Exception:
        pass
    return "n/a"

def load_library(path):
    if not path or not os.path.exists(path):
        raise FileNotFoundError(f"line library file not found: {path}")
    with open(path, "r", encoding="utf-8") as f:
        lib = json.load(f)
    out = []
    for i, e in enumerate(lib):
        try:
            sp = str(e.get("species","")).strip()
            wl = float(e.get("ref_nm"))
            kind = str(e.get("kind","atom")).strip().lower()
        except Exception as ex:
            raise ValueError(f"Invalid library entry at index {i}: {e}") from ex
        out.append(dict(species=sp, ref_nm=wl, kind=("band" if "band" in kind else "atom")))
    return out

def smooth_signal(x,k=7):
    if k<=1: return x.copy()
    k=int(k if k%2==1 else k+1)
    return np.convolve(x, np.ones(k)/k, mode="same")

_SUBS_DIGITS=str.maketrans("0123456789","₀₁₂₃₄₅₆₇₈₉")

def chem_symbol(species):
    if species is None: return ""
    s=str(species).strip()
    repl={"N2 2P":"N2","N2 1P":"N2","N2+ 1N":"N2+","C2 Swan":"C2","CH G-band":"CH","CN violet":"CN","O2 A (abs)":"O2","O2 B (abs)":"O2"}
    for k,v in repl.items():
        if s.startswith(k): s=v; break
    s=s.split()[0]
    s=s.translate(_SUBS_DIGITS).replace("+","⁺").replace("-","⁻")
    return s

def get_font_bold(size):
    for p in ["/usr/share/fonts/truetype/tenorite/Tenorite-Bold.ttf",
              "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
              "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
              "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf"]:
        if os.path.exists(p):
            try: return ImageFont.truetype(p,size=size)
            except Exception: pass
    return ImageFont.load_default()

def build_default_library():
    raise RuntimeError('build_default_library() removed in v6.2; use load_library(library_path)')

# ------------------------------------------------------------------
#  Peak-detection, platåer, lägen
# ------------------------------------------------------------------

def local_maxima_plateau(y, rel_eps=0.01, rel_height=0.03):
    y = np.asarray(y, float); n = len(y); peaks = []
    if n < 3: return np.array([], int)
    ymax = float(np.max(y)); thr = rel_height * ymax
    i = 1
    while i < n - 1:
        if y[i] < thr:
            i += 1; continue
        if not ((y[i] >= y[i-1] and y[i] >= y[i+1]) or (y[i] > y[i-1] and y[i] >= y[i+1])):
            i += 1; continue
        yi = y[i]; eps = rel_eps * max(1e-12, yi)
        cond = max(thr, yi - eps)
        left = i
        while left - 1 >= 0 and y[left - 1] >= cond:
            left -= 1
        right = i
        while right + 1 < n and y[right + 1] >= cond:
            right += 1
        peaks.append((left + right)//2)
        i = right + 1
    return np.array(sorted(set(peaks)), int)

def detect_peaks_multi(nm_axis, I_norm):
    configs = [(9,0.08),(7,0.06),(5,0.04),(3,0.03)]
    peaks = []
    for k, thr in configs:
        y = smooth_signal(I_norm, k=k)
        idx = local_maxima_plateau(y, rel_eps=0.01, rel_height=thr)
        peaks.extend(idx.tolist())
    i_max = int(np.argmax(I_norm))
    peaks.append(i_max)
    peaks = sorted(set(peaks))
    final = []
    for i in peaks:
        if not final:
            final.append(i); continue
        nm_i = nm_axis[i]; nm_prev = nm_axis[final[-1]]
        if abs(nm_i - nm_prev) <= 1.5:
            if I_norm[i] > I_norm[final[-1]]:
                final[-1] = i
        else:
            final.append(i)
    return np.array(final, int)

def estimate_peaks_with_widths(nm_axis, I_norm, peak_idx=None):
    import numpy as _np
    if peak_idx is None:
        peak_idx = detect_peaks_multi(nm_axis, I_norm)
    out = []
    for i in _np.array(peak_idx, int):
        y0 = float(I_norm[i])
        x0 = float(nm_axis[i])
        hm = 0.5 * y0
        L=i
        while L-1>=0 and I_norm[L] >= hm: L-=1
        xL = float(nm_axis[L])
        R=i
        while R+1<len(I_norm) and I_norm[R] >= hm: R+=1
        xR = float(nm_axis[R])
        fwhm = max(0.0, xR - xL)
        out.append(dict(idx=int(i), nm=x0, amp=y0, fwhm=fwhm, prom=y0))
    return out

def detect_mode_auto(peaks, nm_axis, I, tol_nm=2.0):
    """Return ('fluorescent'|'neon'|'mercury'|'unknown', scores)."""
    def narrow_near(target):
        for p in peaks:
            if abs(p['nm']-target) <= tol_nm and p.get('fwhm',10) < 2.5:
                return True
        return False
    has_hg436 = narrow_near(435.833)
    has_hg546 = narrow_near(546.074)
    broad_611 = any(abs(p['nm']-611)<=6 and p.get('fwhm',0)>=5 for p in peaks)
    broad_544 = any(abs(p['nm']-544)<=4 and p.get('fwhm',0)>=5 for p in peaks)
    neon_cluster = sum(1 for p in peaks if 585<=p['nm']<=640 and p.get('fwhm',10)<2.5) >= 3
    scores = {'hg_lines': int(has_hg436)+int(has_hg546),
              'broad_611': int(broad_611), 'broad_544': int(broad_544),
              'neon_cluster': int(neon_cluster)}
    if (has_hg436 and has_hg546) and (broad_611 or broad_544): return 'fluorescent', scores
    if neon_cluster and not (broad_611 or broad_544): return 'neon', scores
    if has_hg436 and has_hg546: return 'mercury', scores
    return 'unknown', scores

# ------------------------------------------------------------------
#  Band-geometri & emission/absorption-hint
# ------------------------------------------------------------------

def estimate_band_row(img):
    im = img if isinstance(img, Image.Image) else Image.open(img)
    H = im.height
    y = H//2
    return y

def band_emission_hint(band_img_path):
    try:
        img = Image.open(band_img_path).convert("L")
        arr = np.array(img, np.float32)
        row = estimate_band_row(img); W = 50
        r0 = max(0, row - W//2); r1 = min(arr.shape[0], row + W//2 + 1)
        roi = arr[r0:r1, :]

        p5  = float(np.percentile(roi, 5))
        med = float(np.median(roi))
        p95 = float(np.percentile(roi, 95))

        up   = max(0.0, p95 - med)
        down = max(0.0, med - p5)

        hi_thr = med + 0.6*up
        lo_thr = max(0.0, med - 0.6*down)
        bright_frac = float(np.mean(roi > hi_thr))
        dark_frac   = float(np.mean(roi < lo_thr))

        if up > 1.8*down and med <= (p95*0.7):
            score = min(1.0, (up/(down+1e-6) - 1.8)/1.5 + 0.5) * (1.0 - min(0.6, bright_frac))
            return "emission", max(0.25, float(score))
        if down > 1.8*up and med >= (p95*0.6):
            score = min(1.0, (down/(up+1e-6) - 1.8)/1.5 + 0.5) * (1.0 - min(0.7, dark_frac))
            return "absorption", max(0.25, float(score))

        if med < 60 and p95 > 140:
            score = min(1.0, (p95 - med)/200.0) * (1.0 - min(0.6, bright_frac)); return "emission", score
        if med > 120 and p5 < 60:
            score = min(1.0, (med - p5)/200.0) * (1.0 - min(0.7, dark_frac));   return "absorption", score

        nz_frac10 = float(np.mean(roi > (p5 + 10.0)))
        if med <= 5 and nz_frac10 < 0.25 and (p95 - med) > 10:
            score = min(1.0, (0.25 - nz_frac10)/0.25) + min(0.5, (p95 - med)/80.0)
            return "emission", max(0.35, float(min(1.0, score)))

        return None, 0.0
    except Exception:
        return None, 0.0

# ------------------------------------------------------------------
#  Matchning, domänregler, gating
# ------------------------------------------------------------------

def initial_match(peaks_nm, peaks_I, lib_atoms, tol_nm):
    matches=[]
    for nm_meas, inten in zip(peaks_nm, peaks_I):
        best=None; best_score=-1.0
        for entry in lib_atoms:
            d=abs(entry["ref_nm"]-nm_meas)
            if d<=tol_nm:
                score=max(0.0,1.0-d/tol_nm)*100.0*float(inten)
                if score>best_score: best=entry; best_score=score
        if best is not None:
            matches.append(dict(
                nm_meas=float(nm_meas),
                corrected_nm=float(nm_meas),
                pixel=None,
                type="em",
                species=best["species"],
                ref_nm=best["ref_nm"],
                delta_nm=float(best["ref_nm"]-nm_meas),
                rel_strength=float(inten),
                score=float(best_score),
                flags=[]
            ))
    return matches

def robust_offset_estimate(peaks_nm, peaks_I, lib_atoms, base_tol):
    tmp=initial_match(peaks_nm,peaks_I,lib_atoms,tol_nm=max(3.0,base_tol*3.0))
    if not tmp: return 0.0
    diffs=np.array([m["ref_nm"]-m["nm_meas"] for m in tmp],float)
    q1,q3=np.percentile(diffs,[25,75]); iqr=q3-q1; lo,hi=q1-1.5*iqr, q3+1.5*iqr
    keep=(diffs>=lo)&(diffs<=hi)
    if keep.sum()>=3: diffs=diffs[keep]
    return float(np.median(diffs))

def apply_domain_rules(matches, tol_nm):
    out = []
    hg_present = {round(m["ref_nm"],3) for m in matches if m.get("species")=="Hg" and abs(m.get("delta_nm",0.0)) <= max(2.5,tol_nm)}
    na_present = {round(m["ref_nm"],3) for m in matches if m.get("species")=="Na"}
    he_sig     = {447.15,471.31,492.19,501.57,706.52}
    he_present = {round(m["ref_nm"],2) for m in matches if m.get("species")=="He"}
    ar_refs    = [m for m in matches if m.get("species")=="Ar" and 696<=m.get("ref_nm",0)<=912.5]
    hg_ok = len(hg_present) >= 1
    na_ok = {588.995,589.592}.issubset(na_present)
    he_ok = len(he_sig.intersection(he_present)) >= 2
    ar_ok = len(ar_refs) >= 2
    for m in matches:
        m = dict(m)
        sp = m.get("species","")
        weak = ((sp=="Hg" and not hg_ok) or
                (sp=="Na" and not na_ok) or
                (sp=="He" and not he_ok) or
                (sp=="Ar" and not ar_ok))
        if weak:
            m.setdefault("flags", [])
            if "weak_evidence" not in m["flags"]:
                m["flags"] = list(m["flags"]) + ["weak_evidence"]
            m["score"] = float(max(0.0, m.get("score",0.0)*0.6))
        out.append(m)
    return out

def gate_matches(matches, nm_min, nm_max):
    if not matches: return []
    matches=sorted(matches,key=lambda m:m.get("score",0.0), reverse=True)
    kept=[]; species_seen=[]
    for m in matches:
        s0=str(m.get("species","")).split()[0]
        if s0 not in species_seen and len(set(species_seen))>=MAX_SPECIES_GLOBAL:
            continue
        kept.append(m); species_seen.append(s0)
    matches=kept
    final=[]
    w_lo=math.floor(nm_min/50)*50; w_hi=math.ceil(nm_max/50)*50
    for w0 in np.arange(w_lo,w_hi,50.0):
        w1=w0+50.0
        bucket=[m for m in matches if w0<=m["ref_nm"]<w1]
        species_in_bucket=[]; bucket_kept=[]
        for m in bucket:
            s0=str(m.get("species","")).split()[0]
            if s0 not in species_in_bucket and len(set(species_in_bucket))>=MAX_SPECIES_PER_50NM:
                continue
            bucket_kept.append(m); species_in_bucket.append(s0)
        bucket_kept=sorted(bucket_kept,key=lambda m:m.get("score",0.0), reverse=True)[:MAX_LABELS_PER_50NM]
        final.extend(bucket_kept)
    final=sorted(final,key=lambda m:m.get("score",0.0), reverse=True)[:MAX_LABELS]
    return final

def sigma_from_scores(matches):
    if not matches: return
    scores=sorted([m.get("score",0.0) for m in matches])
    def level(s):
        p=sum(1 for x in scores if x<=s)/max(1,len(scores))
        if p>=0.90: return 6
        if p>=0.75: return 5
        if p>=0.60: return 4
        if p>=0.40: return 3
        if p>=0.20: return 2
        return 1
    for m in matches:
        m["sigma"]=level(m.get("score",0.0))

def plateau_segments(y, x_axis, rel_top=0.92, rel_eps=0.012):
    y = np.asarray(y, float); n = len(y)
    if n == 0: return []
    ymax = float(np.max(y)); thr = rel_top * ymax
    mask = y >= thr
    segs = []; i = 0
    while i < n:
        if not mask[i]:
            i += 1; continue
        j = i
        while j < n and mask[j]:
            j += 1
        eps = rel_eps * ymax
        L = i
        while L - 1 >= 0 and y[L - 1] >= thr - eps:
            L -= 1
        R = j - 1
        while R + 1 < n and y[R + 1] >= thr - eps:
            R += 1
        nm_lo = float(x_axis[L]); nm_hi = float(x_axis[R]); nm_mid = 0.5 * (nm_lo + nm_hi)
        segs.append((nm_lo, nm_hi, nm_mid, L, R))
        i = R + 1
    return segs

def has_local_peak(nm_axis, I_norm, x_nm, window_nm=1.6):
    idx=np.searchsorted(nm_axis,x_nm)
    L=np.searchsorted(nm_axis,x_nm-window_nm); R=np.searchsorted(nm_axis,x_nm+window_nm)
    L=max(0,L); R=min(len(nm_axis)-1,R)
    if R<=L+1: return False
    seg=I_norm[L:R+1]
    m=float(np.max(seg)); mi=float(np.min(seg)); med=float(np.median(seg)); std=float(np.std(seg))
    irel=int(np.argmax(seg)); interior=(0<irel<len(seg)-1)
    cond_global = m >= 0.05*(np.max(I_norm)+1e-12)
    cond_local  = (m - med) >= max(2.0*std, 0.08) or (m - mi) >= 0.35*(max(1e-9,(np.max(seg)-mi)))
    eps=0.01*m; plateau_span = np.sum(seg >= m-eps) >= 2
    return interior and (cond_global or cond_local or plateau_span)

def post_qc(matches, nm_axis, I_norm, px_axis, band_img_size, qc_log):
    keep=[]; removed=[]
    for m in matches:
        x_check=float(m.get("corrected_nm", m.get("nm_meas", m["ref_nm"])))
        ok_peak=has_local_peak(nm_axis,I_norm,x_check,window_nm=3.2)
        if ok_peak: keep.append(m)
        else: removed.append((m,"no_local_peak"))
    if removed:
        qc_log.append(f"Removed {len(removed)} labels without clear local peak near corrected/measured nm (plateau-aware).")
    keep=gate_matches(keep,float(nm_axis.min()),float(nm_axis.max()))
    return keep

def add_plateau_guesses(matches, offset, nm_axis, I_norm, lib_atoms):
    segs=plateau_segments(I_norm, nm_axis, rel_top=0.92, rel_eps=0.012)
    if not segs: return matches
    existing_positions=[float(m.get("corrected_nm", m.get("nm_meas", m["ref_nm"]))) for m in matches]
    new_matches=[]
    for (nm_lo, nm_hi, nm_mid, L, R) in segs:
        if any(abs(x-nm_mid)<=3.0 for x in existing_positions): continue
        cand=[e for e in lib_atoms if abs(e["ref_nm"]-nm_mid)<=15.0]
        if not cand: continue
        cand=sorted(cand, key=lambda e: abs(e["ref_nm"]-nm_mid))[:2]
        for e in cand:
            nm_meas=nm_mid; corrected=nm_meas+float(offset); delta=float(e["ref_nm"]-corrected)
            inten=float(np.max(I_norm[L:R+1])); score=10.0*max(0.0, 1.0-abs(e["ref_nm"]-nm_mid)/15.0)
            new_matches.append(dict(nm_meas=float(nm_meas), corrected_nm=float(corrected), pixel=None, type="em",
                                    species=e["species"], ref_nm=float(e["ref_nm"]), delta_nm=float(delta),
                                    rel_strength=float(inten), score=float(score), flags=["plateau_guess"]))
    if new_matches:
        seen={(m['species'], round(m['ref_nm'],3)) for m in matches}
        for m in new_matches:
            key=(m['species'], round(m['ref_nm'],3))
            if key not in seen:
                matches.append(m); seen.add(key)
    return matches

# ------------------------------------------------------------------
#  Annotering – band + graf
# ------------------------------------------------------------------

def draw_band_annot(band_path, matches, nm_axis, px_axis, out_path):
    img=Image.open(band_path).convert("RGB"); W,H=img.size; y_band=estimate_band_row(img)
    px_sorted=np.array(px_axis,float); nm_sorted=np.array(nm_axis,float); order=np.argsort(nm_sorted)
    px_sorted=px_sorted[order]; nm_sorted=nm_sorted[order]
    def px_from_nm(nm_val): return float(np.interp(nm_val, nm_sorted, px_sorted))
    px_min=float(px_sorted.min()); px_max=float(px_sorted.max())
    for m in matches:
        nm_use=float(m.get("nm_meas", m.get("corrected_nm", m["ref_nm"])))
        px_calc=px_from_nm(nm_use); m["pixel"]=float(px_calc)
        x_img=(px_calc-px_min)*(W/max(1e-9,(px_max-px_min))); m["x_img"]=float(max(2,min(W-2,x_img)))
    min_dx=max(28,min(34,W//40))
    xs=sorted([(m["x_img"],idx) for idx,m in enumerate(matches)], key=lambda t:t[0])
    tiers={xs[0][1]:0} if xs else {}
    for i in range(1,len(xs)):
        prev_x,prev_idx=xs[i-1]; cur_x,cur_idx=xs[i]
        if cur_x-prev_x<min_dx: tiers[cur_idx]=tiers.get(prev_idx,0)+1
        else: tiers[cur_idx]=0
    draw=ImageDraw.Draw(img); font=get_font_bold(size=max(11,min(20,W//110)))
    for idx,m in enumerate(matches):
        y0=min(H-2,int(y_band+12)); y1=max(0,y0-12)
        draw.line([(m["x_img"],y0),(m["x_img"],y1)], fill=(215,215,215), width=2)
        label=chem_symbol(m["species"]); tw,th=draw.textsize(label,font=font)
        tier=tiers.get(idx,0); block_h=2*th
        tx=max(2,min(W-tw-2,int(m["x_img"]-tw//2)))
        ty=min(H-(tier+1)*block_h-4, y0+2+tier*block_h)
        draw.text((tx+1,ty+1),label,font=font,fill=(0,0,0)); draw.text((tx,ty),label,font=font,fill=(255,255,255))
        sig=f"σ={m.get('sigma',3)}"; sw,sh=draw.textsize(sig,font=font)
        sx=max(2,min(W-sw-2,int(m["x_img"]-sw//2))); sy=ty+th-1
        draw.text((sx+1,sy+1),sig,font=font,fill=(0,0,0)); draw.text((sx,sy),sig,font=font,fill=(255,255,255))
    if W<1200:
        scale=1200.0/W; img=img.resize((1200,int(H*scale)), resample=Image.LANCZOS)
    img.save(out_path)

def draw_graph_annot(nm_axis, I_norm, matches, out_png, mode="emission", offset=0.0):
    fig,ax=plt.subplots(figsize=(11,4.5), dpi=220)
    nm_axis_plot = np.array(nm_axis, float) + float(offset)
    sig = I_norm if mode!="absorption" else (1.0 - I_norm)
    ax.plot(nm_axis_plot, sig, linewidth=2.5, color="#c58b00")
    ax.set_xlabel("Våglängd (nm)")
    ax.set_ylabel("Normaliserad intensitet")
    ax.set_title("Spektrumprofil med identifierade ämnen")
    ax.set_ylim(0,1.2); ax.grid(True, which="both", linestyle="--", linewidth=0.5, alpha=0.5)
    placed=[]
    for m in matches:
        x=float(m.get("corrected_nm", m.get("nm_meas", m["ref_nm"])))
        # v6.2.1 FIX: interpolera mot sig (plottad signal), inte alltid mot emission-I_norm
        y=float(np.interp(x, nm_axis_plot, sig))
        txt=chem_symbol(m["species"]); y_try=min(1.18, y+0.03)
        for _ in range(10):
            bbox=(x-6,y_try-0.05,x+6,y_try+0.05)
            if not any((bbox[2]>b[0] and bbox[0]<b[2] and bbox[3]>b[1] and bbox[1]<b[3]) for b in placed):
                ax.plot([x],[y], marker=(3,0,0), markersize=8, color="k")
                ax.text(x,y_try,txt,ha='center',va='bottom',fontsize=9,fontweight='bold',color='k')
                placed.append(bbox); break
            y_try=min(1.18,y_try+0.035)
    fig.tight_layout(); fig.savefig(out_png, bbox_inches="tight"); plt.close(fig)

# ------------------------------------------------------------------
#  DOCX-generatorer (två varianter)
# ------------------------------------------------------------------

def build_docx_v615(outdir, annotated_band_path, annotated_graph_path, params, matches, peak_count, meta, qc_log, lib_atoms=None, investigation_img=None, original_band_path=None):
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.shared import OxmlElement, qn

    def add_table_borders(table):
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            table._tbl._element.append(tblPr)
        tblBorders = OxmlElement('w:tblBorders')
        for edge in ('top','left','bottom','right','insideH','insideV'):
            elem = OxmlElement(f'w:{edge}')
            elem.set(qn('w:val'),'single'); elem.set(qn('w:sz'),'6'); elem.set(qn('w:color'),'222222')
            tblBorders.append(elem)
        tblPr.append(tblBorders)

    def style_header_row(row):
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), "222222"); tcPr.append(shd)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.color.rgb = RGBColor(255,255,255)
                    run.font.bold = True
                    run.font.size = Pt(8.5)

    def center_all_cells(table):
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    def shrink_table(table, header_pt=8.5, body_pt=10.0):
        if table.rows:
            for cell in table.rows[0].cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(header_pt)
        for row in list(table.rows)[1:]:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(body_pt)

    def _super_sub(txt: str) -> str:
        if not isinstance(txt, str): return txt
        for d, sub in zip("0123456789", "₀₁₂₃₄₅₆₇₈₉"):
            for el in ("H","He","C","N","O","Na","Mg","Al","Si","P","S","Cl","K","Ca","Fe","Li"):
                txt = txt.replace(el + d, el + sub)
        txt = (txt.replace("^-1","⁻¹").replace("^-2","⁻²").replace("^-3","⁻³")
                   .replace("^2","²").replace("^3","³"))
        return txt

    from collections import Counter
    try:
        species_counts = Counter([m.get("species","").split()[0] for m in (matches or [])])
    except Exception:
        species_counts = Counter()

    doc = Document()
    try:
        style = doc.styles['Normal']
        style.font.name = "Tenorite"
        style.font.size = Pt(11)
        pf = style.paragraph_format
        pf.space_before = Pt(0); pf.space_after = Pt(0); pf.line_spacing = 1.0
    except Exception:
        pass

    section = doc.sections[0]
    section.left_margin = Cm(1.7); section.right_margin = Cm(1.7)

    # Sida 1
    p = doc.add_paragraph()
    ver = globals().get("__version__", "")
    run = p.add_run(f"SpectrometerGPT-rapport – SPECTRA {ver}")
    run.bold = True; run.font.size = Pt(20); p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if investigation_img and os.path.exists(investigation_img):
        try:
            doc.add_picture(investigation_img, width=Inches(6.7))
        except Exception:
            pass
        doc.add_paragraph("")
    (_r:=doc.add_paragraph().add_run("ABSTRACT")).bold = True; _r.font.size = Pt(14)
    try:
        abstract_text = build_dynamic_abstract(meta, matches, peak_count, lib_atoms)
    except Exception:
        abstract_text = ""
    paras = [s.strip() for s in str(abstract_text).split("\n\n") if s.strip()]
    if len(paras) < 3:
        try:
            import re as _re
            _phrase = "De högst prioriterade identifieringarna"
            if _phrase in abstract_text:
                _pre, _post = abstract_text.split(_phrase, 1)
                _sentences = _re.split(r"(?<=\.)\s+", _post.strip())
                _mid = _phrase + " " + (_sentences[0] if _sentences else _post.strip())
                if len(_sentences) > 1:
                    _mid += " " + _sentences[1]
                    _rest = " ".join(_sentences[2:])
                else:
                    _rest = ""
                paras = [_pre.strip(), _mid.strip(), _rest.strip()]
                paras = [p for p in paras if p]
            else:
                _words = abstract_text.split()
                if len(_words) >= 30:
                    n = len(_words)
                    paras = [" ".join(_words[:n//3]), " ".join(_words[n//3:2*n//3]), " ".join(_words[2*n//3:])]
                else:
                    paras = [abstract_text]
        except Exception:
            paras = [abstract_text] if abstract_text else []
    paras = paras[:3]
    for i, para in enumerate(paras):
        p_abs = doc.add_paragraph()
        p_abs.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_abs = p_abs.add_run(para)
        if i == 1:
            run_abs.bold = True
        if i < len(paras)-1:
            doc.add_paragraph("")
    doc.add_page_break()

    # Sida 2
    (_h:=doc.add_paragraph().add_run("Bandbild (Original)")).bold = True; _h.font.size = Pt(12)
    if original_band_path and os.path.exists(original_band_path):
        try:
            doc.add_picture(original_band_path, width=Inches(6.7))
            doc.add_paragraph("")
        except Exception: pass
    else:
        try:
            doc.add_picture(annotated_band_path, width=Inches(6.7))
            doc.add_paragraph("")
        except Exception: pass
    (_h:=doc.add_paragraph().add_run("Bandbild (annoterad)")).bold = True; _h.font.size = Pt(12)
    try:
        doc.add_picture(annotated_band_path, width=Inches(6.7))
    except Exception: pass
    doc.add_page_break()

    # Sida 3
    (_h:=doc.add_paragraph().add_run("Spektrumprofil")).bold = True; _h.font.size = Pt(12)
    try:
        doc.add_paragraph("Normaliserad intensitetsprofil – toppar enligt annoterad graf.")
        doc.add_picture(annotated_graph_path, width=Inches(6.7))
    except Exception: pass
    doc.add_paragraph("")

    (_h:=doc.add_paragraph().add_run("Metodtabell")).bold = True; _h.font.size = Pt(12)
    t = doc.add_table(rows=1, cols=4); add_table_borders(t); style_header_row(t.rows[0])
    h = t.rows[0].cells; h[0].text="Parameter"; h[1].text="Värde"; h[2].text="Kommentar"; h[3].text="Enhet"
    rows = [
        ("Smoothing k", str(params.get("k","")), "Rörligt medel (udda)", "-"),
        ("Tröskel", f"{int(params.get('threshold',0.0)*100)} %", "Relativ till max", "%"),
        ("Min avstånd", f"{params.get('min_dist_nm','')}", "Mellan toppar", "nm"),
        ("Tolerans", f"±{params.get('tolerance_nm','')}", "Matchning & rematch", "nm"),
        ("I-def.", "max(R,G,B)", "Auto avgör emission/absorption", "-"),
    ]
    for row in rows:
        r = t.add_row().cells
        r[0].text, r[1].text, r[2].text, r[3].text = row
    shrink_table(t)
    doc.add_paragraph("")

    (_h:=doc.add_paragraph().add_run("Workflow (pipeline)")).bold = True; _h.font.size = Pt(12)
    for _line in [
        "Inläsning av dataset + bilder.",
        "Kalibrering px→nm (datasetets nm-kolumn, annars fabriksfit).",
        "Normalisering & smoothing.",
        "Auto-läge: emission + absorption → bästa helhet.",
        "Grov match mot linjebibliotek + robust offset-korrektion.",
        "Rematch inom tolerans + domänregler.",
        "Gating per 50 nm + artbegränsningar.",
        "Annotering band + graf.",
        "Post-QC + export (PNG/CSV/DOCX).",
    ]:
        doc.add_paragraph("• " + _line)

    doc.add_paragraph().add_run("Analysatorns resonemang (sammanfattning)").bold = True
    doc.add_paragraph("Resultaten baseras på stabil px→nm-relation; felkällor: mättnad, låg SNR, blendade linjer, RGB-viktning. QC och domänregler prioriterar fysiskt rimliga signaturer.")
    doc.add_page_break()

    # Sida 4
    (_h:=doc.add_paragraph().add_run("Instrument/kalibrering")).bold = True; _h.font.size = Pt(12)
    doc.add_paragraph("Instrument: KVANT SPECTRA 1 – synligt ljus (~360–930 nm); <1.8 nm FWHM; <0.5 nm/px; gitter 500 linjer/mm; 1/3″ färg-CMOS 1.3 MP (1280×960).")
    doc.add_paragraph("Fabrikskalibrering (px→nm):")
    tcal = doc.add_table(rows=1, cols=2); add_table_borders(tcal); style_header_row(tcal.rows[0])
    tcal.rows[0].cells[0].text = "Pixel (px)"; tcal.rows[0].cells[1].text = "Våglängd (nm)"
    for _px, _nm in [(32,388.86),(515,587.57),(1110,837.76)]:
        _r = tcal.add_row().cells; _r[0].text=str(_px); _r[1].text=f"{_nm:.2f}"
    doc.add_paragraph("Fabriksfit: nm = a2·px² + a1·px + a0; a2=8.457e−06, a1=0.406760986, a0=375.834988.")

    (_h:=doc.add_paragraph().add_run("Primära indikatorer (σ)")).bold = True; _h.font.size = Pt(12)
    t2 = doc.add_table(rows=1, cols=6); add_table_borders(t2); style_header_row(t2.rows[0])
    hh = t2.rows[0].cells; hh[0].text="mät-nm"; hh[1].text="ref-nm"; hh[2].text="|Δλ|"; hh[3].text="art"; hh[4].text="orsak"; hh[5].text="σ"
    for m in sorted(matches or [], key=lambda z:z.get("score",0.0), reverse=True)[:max(5, len(matches or []))]:
        rr = t2.add_row().cells
        rr[0].text=f"{m.get('corrected_nm', m.get('nm_meas',0.0)):.2f}"
        rr[1].text=f"{m.get('ref_nm',0.0):.2f}"
        rr[2].text=f"{abs(m.get('delta_nm', m.get('ref_nm',0.0)-m.get('nm_meas',0.0))):.2f}"
        rr[3].text=_super_sub(m.get('species',''))
        rr[4].text="proximity/intensitet/domänregel"; rr[5].text=str(m.get('sigma',3))
    shrink_table(t2)

    doc.add_paragraph().add_run("Full linjetabell (nm-sorterad)").bold = True
    t3 = doc.add_table(rows=1, cols=10); add_table_borders(t3); style_header_row(t3.rows[0])
    c = t3.rows[0].cells
    c[0].text="Measured_nm"; c[1].text="Corrected_nm"; c[2].text="Pixel"; c[3].text="Type"; c[4].text="Element/Species"; c[5].text="Ref_nm"; c[6].text="Δλ"; c[7].text="Rel.str."; c[8].text="σ"; c[9].text="Källa"
    def _src_of(m):
        t = str(m.get("type",""))
        if "·both" in t: return "Båda"
        if "·band" in t: return "Band"
        if "·xlsx" in t: return "Data"
        return ""
    for m in sorted(matches or [], key=lambda z: z.get("ref_nm", 0.0)):
        r = t3.add_row().cells
        r[0].text=f"{m.get('nm_meas',float('nan')):.3f}"
        r[1].text=f"{m.get('corrected_nm', m.get('nm_meas',float('nan'))):.3f}"
        r[2].text=f"{m.get('pixel', float('nan')):.1f}"
        r[3].text=str(m.get('type','em'))
        r[4].text=_super_sub(m.get('species',''))
        r[5].text=f"{m.get('ref_nm',0.0):.3f}"
        r[6].text=f"{m.get('delta_nm', m.get('ref_nm',0.0)-m.get('nm_meas',0.0)):+.3f}"
        r[7].text=f"{m.get('rel_strength',0.0):.3f}"
        r[8].text=str(m.get('sigma',3))
        r[9].text=_src_of(m)
    center_all_cells(t3); shrink_table(t3)
    doc.add_page_break()

    # Sida 5 – referensbibliotek
    if lib_atoms:
        (_h:=doc.add_paragraph().add_run("Referensbibliotek")).bold = True; _h.font.size = Pt(12)
        entries = sorted(lib_atoms, key=lambda z: z["ref_nm"])
        blocks = 3; cols = blocks * 3
        tlib = doc.add_table(rows=1, cols=cols); add_table_borders(tlib); style_header_row(tlib.rows[0])
        head = ["Art","Ref_nm","Kem-symbol"] * blocks
        for i,txt in enumerate(head): tlib.rows[0].cells[i].text = txt
        rows_needed = (len(entries)+blocks-1)//blocks
        for r_idx in range(rows_needed):
            row_cells = tlib.add_row().cells
            for b in range(blocks):
                idx = r_idx + b*rows_needed
                cbase = b*3
                if idx < len(entries):
                    e = entries[idx]
                    row_cells[cbase+0].text = e["species"]
                    row_cells[cbase+1].text = f"{e['ref_nm']:.3f}"
                    row_cells[cbase+2].text = chem_symbol(e["species"])
                else:
                    row_cells[cbase+0].text = ""
                    row_cells[cbase+1].text = ""
                    row_cells[cbase+2].text = ""
        shrink_table(tlib)
    doc.add_page_break()

    # Sida 6+ – QC + reproducibility
    (_h:=doc.add_paragraph().add_run("Kvalitetsrapport")).bold = True; _h.font.size = Pt(12)
    if qc_log:
        for entry in qc_log: doc.add_paragraph(f"• {entry}")
    else:
        doc.add_paragraph("• OK (inga ändringar)")

    (_h:=doc.add_paragraph().add_run("Reproducerbarhet")).bold = True; _h.font.size = Pt(12)
    doc.add_paragraph(f"Tidsstämpel: {__import__('datetime').datetime.now().isoformat(timespec='seconds')}")
    for k,v in (meta or {}).items():
        doc.add_paragraph(f"{k}: {v}")

    doc.add_paragraph().add_run("Analyslogg (detaljerad)").bold = True
    try:
        atext = params.get("analysis_log","")
    except Exception:
        atext = ""
    if atext:
        for para in str(atext).split("\n"):
            doc.add_paragraph(para)
    else:
        doc.add_paragraph("Ingen detaljerad logg tillgänglig i denna körning.")

    out_docx=os.path.join(outdir,"report.docx")
    doc.save(out_docx)
    return out_docx

def build_docx(outdir, annotated_band_path, annotated_graph_path, params, matches, peak_count, meta, qc_log, lib_atoms=None):
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.shared import OxmlElement, qn

    def add_table_borders(table):
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            table._tbl._element.append(tblPr)
        tblBorders = OxmlElement('w:tblBorders')
        for edge in ('top','left','bottom','right','insideH','insideV'):
            elem = OxmlElement(f'w:{edge}')
            elem.set(qn('w:val'),'single'); elem.set(qn('w:sz'),'6'); elem.set(qn('w:color'),'222222')
            tblBorders.append(elem)
        tblPr.append(tblBorders)

    def style_header_row(row):
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), "222222"); tcPr.append(shd)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.color.rgb = RGBColor(255,255,255)
                    run.font.bold = True
                    run.font.size = Pt(8.5)

    def center_all_cells(table):
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    def shrink_table(table, header_pt=8.5, body_pt=10.0):
        if table.rows:
            for cell in table.rows[0].cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(header_pt)
        for row in list(table.rows)[1:]:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(body_pt)

    def _super_sub(txt: str) -> str:
        if not isinstance(txt, str): return txt
        for d, sub in zip("0123456789", "₀₁₂₃₄₅₆₇₈₉"):
            for el in ("H","He","C","N","O","Na","Mg","Al","Si","P","S","Cl","K","Ca","Fe","Li"):
                txt = txt.replace(el + d, el + sub)
        txt = (txt.replace("^-1","⁻¹").replace("^-2","⁻²").replace("^-3","⁻³")
                   .replace("^2","²").replace("^3","³"))
        return txt

    from collections import Counter
    species_counts = Counter([m["species"].split()[0] for m in matches])
    top_species = ", ".join([f"{s} (n={c})" for s,c in species_counts.most_common()]) if species_counts else "—"

    doc = Document()
    try:
        style = doc.styles['Normal']
        style.font.name = "Tenorite"
        style.font.size = Pt(11)
        pf = style.paragraph_format
        pf.space_before = Pt(0); pf.space_after = Pt(0); pf.line_spacing = 1.0
    except Exception:
        pass

    section = doc.sections[0]
    section.left_margin = Cm(1.7); section.right_margin = Cm(1.7)

    p = doc.add_paragraph()
    run = p.add_run(f"SpectrometerGPT-rapport – SPECTRA {__version__}")
    run.bold = True; run.font.size = Pt(20); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    (_r:=doc.add_paragraph().add_run("ABSTRACT")).bold = True; _r.font.size = Pt(14)

    abstract_text = build_dynamic_abstract(meta, matches, peak_count, lib_atoms)
    paras = [s.strip() for s in str(abstract_text).split("\n\n") if s.strip()]
    if len(paras) < 3:
        try:
            import re as _re
            _phrase = "De högst prioriterade identifieringarna"
            if _phrase in abstract_text:
                _pre, _post = abstract_text.split(_phrase, 1)
                _sentences = _re.split(r"(?<=\.)\s+", _post.strip())
                _mid = _phrase + " " + (_sentences[0] if _sentences else _post.strip())
                if len(_sentences) > 1:
                    _mid += " " + _sentences[1]
                    _rest = " ".join(_sentences[2:])
                else:
                    _rest = ""
                paras = [_pre.strip(), _mid.strip(), _rest.strip()]
                paras = [p for p in paras if p]
            else:
                _words = abstract_text.split()
                if len(_words) >= 30:
                    n = len(_words)
                    paras = [" ".join(_words[:n//3]), " ".join(_words[n//3:2*n//3]), " ".join(_words[2*n//3:])]
                else:
                    paras = [abstract_text]
        except Exception:
            paras = [abstract_text] if abstract_text else []
    paras = paras[:3]
    for i, para in enumerate(paras):
        p_abs = doc.add_paragraph()
        p_abs.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_abs = p_abs.add_run(para)
        if i == 1:
            run_abs.bold = True
        if i < len(paras)-1:
            doc.add_paragraph("")

    doc.add_paragraph().add_run("Bandbild (identifierade linjer/band)").bold = True
    try:
        doc.add_picture(annotated_band_path, width=Inches(6.7))
    except Exception:
        pass
    doc.add_page_break()

    doc.add_paragraph().add_run("Spektrumprofil (identifierade linjer)").bold = True
    try:
        _tops = sorted(matches, key=lambda z: z.get("score",0.0), reverse=True)[:3]
        _txt = ", ".join([f"{chem_symbol(m['species'])} {m['ref_nm']:.1f} nm" for m in _tops]) if _tops else "—"
        doc.add_paragraph("Normaliserad intensitetsprofil med märkta referenslinjer; tydlig(e) topp(ar): " + _txt)
    except Exception:
        doc.add_paragraph("Normaliserad intensitetsprofil med märkta referenslinjer; tydlig(e) topp(ar): —")
    try:
        doc.add_picture(annotated_graph_path, width=Inches(6.7))
    except Exception:
        pass

    doc.add_paragraph().add_run("Metodtabeller").bold = True
    t = doc.add_table(rows=1, cols=4); add_table_borders(t); style_header_row(t.rows[0])
    h = t.rows[0].cells; h[0].text="Parameter"; h[1].text="Värde"; h[2].text="Kommentar"; h[3].text="Enhet"
    rows = [
        ("Smoothing k", str(params["k"]), "Rörligt medel (udda)", "-"),
        ("Tröskel", f"{int(params['threshold']*100)} %", "Relativ till max", "%"),
        ("Min avstånd", f"{params['min_dist_nm']}", "Mellan toppar", "nm"),
        ("Tolerans", f"±{params['tolerance_nm']}", "Matchning & rematch", "nm"),
        ("I-def.", "max(R,G,B)", "Auto väljer emission/absorption", "-"),
    ]
    for row in rows:
        r = t.add_row().cells
        r[0].text, r[1].text, r[2].text, r[3].text = row
    shrink_table(t)

    p_wf = doc.add_paragraph(); r_wf = p_wf.add_run("Workflow (pipeline)"); r_wf.bold = True
    for _line in [
        "Inläsning av dataset (px, R,G,B [+ ev. nm]) och bilder (band, graf).",
        "Kalibrering px→nm: datasetets nm-kolumn; annars fabriksfit (kvadratisk).",
        "Normalisering I=max(R,G,B) och smoothing.",
        "Auto-läge: kör emissions- och absorptionspipeline; välj den med högst matchscore/antal.",
        "Grov matchning mot linjebibliotek; robust median-offset (IQR-trim).",
        "Rematchning inom tolerans; domänregler (Hg/Na/He/Ar).",
        "Gating per 50 nm + global artbegränsning.",
        "Band-annotering och spektrumprofil med små etiketter (ingen skugga).",
        "Post-QC: lokal topp/platå-krav, strykningar och omritning.",
        "Export av annotated_band.png, annotated_graph.png, lines.csv, report.docx.",
    ]:
        doc.add_paragraph("• " + _line)

    try:
        doc.add_paragraph(f"Kandidattoppar: {peak_count}; etiketter efter regler/gating/post-QC: {len(matches)}.")
    except Exception:
        doc.add_paragraph("Kandidattoppar: —; etiketter efter regler/gating/post-QC: —.")

    p_res = doc.add_paragraph(); r_res = p_res.add_run("Analysatorns resonemang (sammanfattning)"); r_res.bold = True
    doc.add_paragraph("Resultaten tolkas under antagande om stabil px→nm-relation. Möjliga felkällor: mättnad i starka linjer, låg SNR, band-blend i klusterregioner samt RGB-viktning. Lokala topp/platå-kravet i korstestet reducerar risken för felmärkning vid platåer. Domänregler och gating prioriterar fysiskt rimliga signaturer framför enstaka slumpträffar.")

    p_ins = doc.add_paragraph(); r_ins = p_ins.add_run("Instrument/kalibrering"); r_ins.bold = True
    doc.add_paragraph("Instrument: KVANT SPECTRA 1 – undervisningsspektrometer för synligt ljus (~360–930 nm). Upplösning <1.8 nm FWHM; pixelupplösning <0.5 nm; gitter 500 linjer/mm; detektor 1/3″ färg-CMOS 1.3 MP (1280×960); mjukvara för Windows 8/10/11; kapslad, robust konstruktion för elevlabb. Källa: forschool.eu / KVANT.")
    doc.add_paragraph("Fabrikskalibrering (px→nm):")
    tcal = doc.add_table(rows=1, cols=2); add_table_borders(tcal); style_header_row(tcal.rows[0])
    tcal.rows[0].cells[0].text = "Pixel (px)"; tcal.rows[0].cells[1].text = "Våglängd (nm)"
    for _px, _nm in [(32,388.86),(515,587.57),(1110,837.76)]:
        _r = tcal.add_row().cells; _r[0].text=str(_px); _r[1].text=f"{_nm:.2f}"
    doc.add_paragraph("Fabriksfit (px→nm): nm = a2·px² + a1·px + a0; a2=8.457e−06, a1=0.406760986, a0=375.834988. Riktning nm L→R.")
    try:
        _cal = str(meta.get("Kalibrering",""))
        if "dataset" in _cal:
            doc.add_paragraph("Kalibrering: datasetets nm-kolumn användes direkt.")
        else:
            doc.add_paragraph("Kalibrering: fabriksfit användes (kvadratisk px→nm).")
    except Exception:
        doc.add_paragraph("Kalibrering: —")

    doc.add_paragraph().add_run("Primära indikationer (σ)").bold = True
    t2 = doc.add_table(rows=1, cols=6); add_table_borders(t2); style_header_row(t2.rows[0])
    hh = t2.rows[0].cells; hh[0].text="mät-nm"; hh[1].text="ref-nm"; hh[2].text="|Δλ|"; hh[3].text="art"; hh[4].text="orsak"; hh[5].text="σ"
    for m in sorted(matches, key=lambda z:z.get("score",0.0), reverse=True)[:max(5, len(matches))]:
        rr = t2.add_row().cells
        rr[0].text=f"{m.get('corrected_nm', m.get('nm_meas',0.0)):.2f}"
        rr[1].text=f"{m['ref_nm']:.2f}"
        rr[2].text=f"{abs(m.get('delta_nm', m['ref_nm']-m.get('nm_meas',0.0))):.2f}"
        rr[3].text=_super_sub(m['species'])
        rr[4].text="proximity/intensitet/domänregel"; rr[5].text=str(m.get('sigma',3))
    shrink_table(t2)

    doc.add_paragraph().add_run("Full linjetabell (nm-sorterad)").bold = True
    t3 = doc.add_table(rows=1, cols=10); add_table_borders(t3); style_header_row(t3.rows[0])
    c = t3.rows[0].cells
    c[0].text="Measured_nm"; c[1].text="Corrected_nm"; c[2].text="Pixel"; c[3].text="Type"; c[4].text="Element/Species"; c[5].text="Ref_nm"; c[6].text="Δλ"; c[7].text="Rel.str."; c[8].text="σ"; c[9].text="Källa"

    def _src_of(m):
        t = str(m.get("type",""))
        if "·both" in t: return "Båda"
        if "·band" in t: return "Band"
        if "·xlsx" in t: return "Data"
        return ""

    for m in sorted(matches, key=lambda z: z["ref_nm"]):
        r = t3.add_row().cells
        r[0].text=f"{m.get('nm_meas',float('nan')):.3f}"
        r[1].text=f"{m.get('corrected_nm', m.get('nm_meas',float('nan'))):.3f}"
        r[2].text=f"{m.get('pixel', float('nan')):.1f}"
        r[3].text=m.get("type","em")
        r[4].text=_super_sub(m["species"])
        r[5].text=f"{m['ref_nm']:.3f}"
        r[6].text=f"{m.get('delta_nm', m['ref_nm']-m.get('nm_meas',0.0)):+.3f}"
        r[7].text=f"{m.get('rel_strength',0.0):.3f}"
        r[8].text=str(m.get('sigma',3))
        r[9].text=_src_of(m)

    center_all_cells(t3)
    shrink_table(t3)

    if lib_atoms:
        doc.add_paragraph().add_run("Bilaga: Referensbibliotek (kompakt tabell)").bold = True
        entries = sorted(lib_atoms, key=lambda z: z["ref_nm"])
        blocks = 3; cols = blocks * 3
        tlib = doc.add_table(rows=1, cols=cols); add_table_borders(tlib); style_header_row(tlib.rows[0])
        head = ["Art","Ref_nm","Kem-symbol"] * blocks
        for i,txt in enumerate(head): tlib.rows[0].cells[i].text = txt
        rows_needed = (len(entries)+blocks-1)//blocks
        for r_idx in range(rows_needed):
            row_cells = tlib.add_row().cells
            for b in range(blocks):
                idx = r_idx + b*rows_needed
                cbase = b*3
                if idx < len(entries):
                    e = entries[idx]
                    row_cells[cbase+0].text = e["species"]
                    row_cells[cbase+1].text = f"{e['ref_nm']:.3f}"
                    row_cells[cbase+2].text = chem_symbol(e["species"])
                else:
                    row_cells[cbase+0].text = ""
                    row_cells[cbase+1].text = ""
                    row_cells[cbase+2].text = ""
        shrink_table(tlib)

    doc.add_paragraph().add_run("Kvalitetsrapport").bold = True
    if qc_log:
        for entry in qc_log:
            doc.add_paragraph(f"• {entry}")
    else:
        doc.add_paragraph("• OK (inga ändringar)")

    doc.add_paragraph().add_run("Reproducerbarhet").bold = True
    doc.add_paragraph(f"Tidsstämpel: {__import__('datetime').datetime.now().isoformat(timespec='seconds')}")
    for k,v in meta.items():
        doc.add_paragraph(f"{k}: {v}")

    doc.add_page_break()
    doc.add_paragraph().add_run("Analyslogg (detaljerad)").bold = True
    atext = params.get("analysis_log","")
    if atext:
        for para in str(atext).split("\n"):
            doc.add_paragraph(para)
    else:
        doc.add_paragraph("Ingen detaljerad logg tillgänglig i denna körning.")

    out_docx=os.path.join(outdir,"report.docx")
    doc.save(out_docx)
    return out_docx

# ------------------------------------------------------------------
#  Kärn-engine (v6.2, patchad -> v6.2.1)
# ------------------------------------------------------------------

def _eval_pipeline(nm, I_norm, px, lib_atoms, min_dist_nm, tolerance_nm):
    idx=detect_peaks_multi(nm,I_norm)
    cand_sorted=sorted(idx.tolist(), key=lambda i:I_norm[i], reverse=True)
    keep_idx=[]
    for i in cand_sorted:
        nm_i=nm[i]
        if all(abs(nm_i-nm[j])>=min_dist_nm for j in keep_idx):
            keep_idx.append(i)
    keep_idx=sorted(keep_idx)
    peaks_nm=nm[keep_idx]; peaks_I=I_norm[keep_idx]
    matches_init=initial_match(peaks_nm,peaks_I,lib_atoms,tol_nm=max(3.0,tolerance_nm))
    offset=robust_offset_estimate(peaks_nm,peaks_I,lib_atoms,base_tol=tolerance_nm)
    corr_peaks_nm=peaks_nm+offset
    matches_corr=initial_match(corr_peaks_nm,peaks_I,lib_atoms,tol_nm=max(2.5,tolerance_nm))
    seen=set(); matches=[]
    for mm in (matches_corr+matches_init):
        key=(mm['species'], round(mm['ref_nm'],3))
        if key not in seen:
            seen.add(key); matches.append(mm)
    for m in matches:
        m["corrected_nm"]=float(m["nm_meas"])+offset
        m["delta_nm"]=float(m["ref_nm"]-m["corrected_nm"])
    matches=apply_domain_rules(matches, tol_nm=tolerance_nm)
    nm_min,nm_max=float(nm.min()),float(nm.max())
    matches=gate_matches(matches,nm_min,nm_max); sigma_from_scores(matches)
    total_score=float(sum(m.get("score",0.0) for m in matches))
    return matches, offset, keep_idx, total_score

def base_run_engine(dataset, band, graph_img_unused, mode="auto", k=7, threshold=0.06, min_dist_nm=2.0, tolerance_nm=2.0, outdir=".", library_path=None):
    os.makedirs(outdir, exist_ok=True); qc_log=[]
    ext=os.path.splitext(dataset)[1].lower()
    if ext in (".xlsx",".xls"):
        xls=pd.ExcelFile(dataset); df=pd.read_excel(dataset, sheet_name=xls.sheet_names[0])
    else:
        df=pd.read_csv(dataset)

    def normcol(c):
        c=str(c).strip().lower()
        c=c.replace("å","a").replace("ä","a").replace("ö","o")
        c=c.replace(" ","").replace("-","").replace("_","")
        return c

    colmap={normcol(c):c for c in df.columns}
    px_col=next((colmap[c] for c in colmap if c in ("px","pixel","pix","x")), None)
    nm_col=next((colmap[c] for c in colmap if c in ("nm","nanometer","wavelen","wavelength")), None)
    r_col=next((colmap[c] for c in colmap if c in ("r","red")), None)
    g_col=next((colmap[c] for c in colmap if c in ("g","green")), None)
    b_col=next((colmap[c] for c in colmap if c in ("b","blue")), None)
    if px_col is None: raise RuntimeError("Pixelkolumn saknas (px/pixel/pix/x).")
    if r_col is None or g_col is None or b_col is None:
        num_cols=[c for c in df.columns if pd.api.types.is_numeric_dtype(df[c])]
        ex=set([col for col in [px_col, nm_col] if col is not None])
        cand=[c for c in num_cols if c not in ex]
        if len(cand)>=3: r_col,g_col,b_col=cand[:3]
        else: raise RuntimeError("Kunde inte hitta R,G,B-kolumner.")

    px=df[px_col].to_numpy(float); R=df[r_col].to_numpy(float); G=df[g_col].to_numpy(float); B=df[b_col].to_numpy(float)
    calib_used="dataset_nm" if nm_col is not None else "factory"
    if nm_col is not None: nm=df[nm_col].to_numpy(float)
    else: nm=np.polyval(np.array([A2,A1,A0]), px)
    if len(nm)>1 and nm[1]<nm[0]:
        nm=nm[::-1]; px=px[::-1]; R=R[::-1]; G=G[::-1]; B=B[::-1]
    I_em=np.max(np.vstack([R,G,B]), axis=0).astype(float)
    I_norm=I_em/(np.max(I_em)+1e-12)

    lib_atoms=load_library(library_path)

    chosen_mode=mode
    mode_guess = "unknown"  # v6.2.1 FIX: spara auto-läge-gissning

    if mode=="auto":
        _peaks = estimate_peaks_with_widths(nm, I_norm, detect_peaks_multi(nm, I_norm))
        mode_guess, _scores = detect_mode_auto(_peaks, nm, I_norm, tol_nm=tolerance_nm)

        hint,hint_score=band_emission_hint(band)
        matches_em,off_em,keep_em,score_em=_eval_pipeline(nm,I_norm,px,lib_atoms,min_dist_nm,tolerance_nm)
        I_abs_norm=(np.max(I_em)-I_em)/(np.max(I_em)+1e-12)
        matches_abs,off_abs,keep_abs,score_abs=_eval_pipeline(nm,I_abs_norm,px,lib_atoms,min_dist_nm,tolerance_nm)

        cand_em=(len(matches_em),score_em,"emission",matches_em,off_em,keep_em,I_norm)
        cand_abs=(len(matches_abs),score_abs,"absorption",matches_abs,off_abs,keep_abs,I_abs_norm)
        choose=cand_em
        if hint=="emission" and hint_score>=0.25:
            if not (cand_abs[0]>=cand_em[0]+5 and cand_abs[1]>cand_em[1]*1.15):
                choose=cand_em
        elif hint=="absorption" and hint_score>=0.25:
            if not (cand_em[0]>=cand_abs[0]+5 and cand_em[1]>cand_abs[1]*1.15):
                choose=cand_abs
        else:
            if (cand_abs[0]>=cand_em[0]+3 and cand_abs[1]>cand_em[1]*1.05):
                choose=cand_abs
            else:
                choose=cand_em
        chosen_mode=choose[2]; matches,offset,keep_idx,I_plot=choose[3],choose[4],choose[5],choose[6]

        if mode_guess == 'fluorescent':
            chosen_mode='emission'
            matches,offset,keep_idx,_=_eval_pipeline(nm,I_norm,px,lib_atoms,min_dist_nm,tolerance_nm)
            I_plot=I_norm
        if chosen_mode=="absorption":
            _hint,_hs=band_emission_hint(band)
            if _hint=="emission" and _hs>=0.20:
                chosen_mode="emission"
                matches,offset,keep_idx,_=_eval_pipeline(nm,I_norm,px,lib_atoms,min_dist_nm,tolerance_nm)
                I_plot=I_norm
    else:
        if mode=="absorption": I_plot=(np.max(I_em)-I_em)/(np.max(I_em)+1e-12)
        else: I_plot=I_norm
        matches,offset,keep_idx,_=_eval_pipeline(nm,I_plot,px,lib_atoms,min_dist_nm,tolerance_nm)

    before=len(matches)
    matches=add_plateau_guesses(matches, offset, nm_axis=nm, I_norm=I_norm, lib_atoms=lib_atoms)
    if len(matches)>before:
        sigma_from_scores(matches)
    if len(matches)<3:
        tol2=max(tolerance_nm*2.0,4.0)
        matches2,offset2,keep2,_=_eval_pipeline(nm,I_norm,px,lib_atoms,min_dist_nm,tol2)
        if len(matches2)>len(matches):
            matches,offset,keep_idx=matches2,offset2,keep2; tolerance_nm=tol2

    band_img=Image.open(band).convert("RGB"); W,H=band_img.size
    matches=post_qc(matches,nm_axis=nm,I_norm=I_norm,px_axis=px,band_img_size=(W,H),qc_log=qc_log)
    out_band=os.path.join(outdir,"annotated_band.png"); out_graph=os.path.join(outdir,"annotated_graph.png")
    draw_band_annot(band,matches,nm_axis=nm,px_axis=px,out_path=out_band)
    draw_graph_annot(nm_axis=nm,I_norm=I_norm,matches=matches,out_png=out_graph,mode=chosen_mode, offset=offset)
    out_csv=os.path.join(outdir,"lines.csv")
    with open(out_csv,"w",newline="",encoding="utf-8") as f:
        w=csv.writer(f)
        w.writerow(["Measured_nm","Corrected_nm","Pixel","Type","Element/Species","Ref_nm","Delta_nm","Relative_Strength","Score","Flags","Sigma"])
        for m in sorted(matches,key=lambda z: z["ref_nm"]):
            w.writerow([
                f"{m['nm_meas']:.3f}",
                f"{m.get('corrected_nm', m['nm_meas']):.3f}",
                f"{m.get('pixel', float('nan')):.1f}",
                m.get("type","em"),
                m["species"],
                f"{m['ref_nm']:.3f}",
                f"{m.get('delta_nm', m['ref_nm']-m['nm_meas']):+.3f}",
                f"{m.get('rel_strength',0.0):.3f}",
                f"{m.get('score',0.0):.1f}",
                "|".join(m.get('flags',[])),
                m.get('sigma',3)
            ])

    analysis_log = build_analysis_log(nm_axis=nm, I_norm=I_norm, peak_idx=keep_idx, matches=matches, offset=offset, qc_log=qc_log, mode_label=chosen_mode)
    params=dict(k=DEFAULT_K, threshold=DEFAULT_THRESHOLD, min_dist_nm=DEFAULT_MIN_DIST_NM, tolerance_nm=tolerance_nm, analysis_log=analysis_log)

    meta={
        "Dataset": (os.path.basename(dataset) if dataset else "n/a") + f" SHA-256: {sha256_safe(dataset)}",
        "Bandbild": (os.path.basename(band) if band else "n/a") + f" SHA-256: {sha256_safe(band)}",
        "Grafbild": (os.path.basename(graph_img_unused) if graph_img_unused else "n/a") + f" SHA-256: {sha256_safe(graph_img_unused)}",
        "Kalibrering": calib_used,
        "Offset (median Δλ)": f"{offset:+.3f} nm",
        "Spektrumläge": chosen_mode,
        "Auto-läge gissning": mode_guess,
        "Version": f"SpectrometerGPT {__version__}"
    }

    out_docx=build_docx(outdir,out_band,out_graph,params,matches,peak_count=len(keep_idx), meta=meta, qc_log=[], lib_atoms=load_library(library_path))
    print("[OK] Outputs ready:")
    print(" ", out_band)
    print(" ", out_graph)
    print(" ", out_csv)
    print(" ", out_docx)

# ------------------------------------------------------------------
#  CLI (”klassisk” engine-körning, med --flaggar)
# ------------------------------------------------------------------

if __name__=="__main__" and _ALLOW_CLI():
    ap=argparse.ArgumentParser(description=f"SpectrometerGPT engine {__version__}")
    ap.add_argument("--dataset", required=True)
    ap.add_argument("--band", required=True)
    ap.add_argument("--graph", required=True)
    ap.add_argument("--outdir", default=".")
    ap.add_argument("--library_path", default="/mnt/data/line_library_v1.json")
    ap.add_argument("--mode", default="auto")
    ap.add_argument("--k", type=float, default=7)
    ap.add_argument("--threshold", type=float, default=0.06)
    ap.add_argument("--min_dist_nm", type=float, default=2.0)
    ap.add_argument("--tolerance_nm", type=float, default=2.0)
    args=ap.parse_args()
    base_run_engine(args.dataset, args.band, args.graph,
                    mode=args.mode, k=args.k, threshold=args.threshold,
                    min_dist_nm=args.min_dist_nm, tolerance_nm=args.tolerance_nm,
                    outdir=args.outdir, library_path=args.library_path)

# ------------------------------------------------------------------
#  engine6 union wrapper (dataset + band) – v6.2.1
# ------------------------------------------------------------------

def _build_dataset_from_band(band_img_path, save_path):
    img = Image.open(band_img_path).convert("RGB")
    arr = np.array(img, dtype=np.float32)
    y = int(np.argmax(arr.mean(axis=(1,2))))
    y0,y1 = max(0, y-20), min(arr.shape[0], y+21)
    roi = arr[y0:y1,:,:]
    R = np.percentile(roi[:,:,0],95,axis=0).astype(np.float32)
    G = np.percentile(roi[:,:,1],95,axis=0).astype(np.float32)
    B = np.percentile(roi[:,:,2],95,axis=0).astype(np.float32)
    px = np.arange(R.shape[0], dtype=np.float32)
    pd.DataFrame({"px":px,"R":R,"G":G,"B":B}).to_csv(save_path, index=False)
    return save_path

def _key_row(r):
    try:
        return f"{str(r['Element/Species']).strip()}_{float(r['Ref_nm']):.3f}"
    except Exception:
        return None

def _union_lines(df_xlsx, df_band):
    if df_xlsx is None or df_xlsx.empty:
        df_x = pd.DataFrame()
    else:
        df_x = df_xlsx.copy(); df_x["__src"]="xlsx"
    if df_band is None or df_band.empty:
        df_b = pd.DataFrame()
    else:
        df_b = df_band.copy(); df_b["__src"]="band"
    all_df = pd.concat([df_x, df_b], ignore_index=True)
    if all_df.empty: return pd.DataFrame()
    for c in ["Score","Sigma","Ref_nm"]:
        if c in all_df.columns:
            all_df[c] = pd.to_numeric(all_df[c], errors="coerce")
    all_df["__key"] = all_df.apply(_key_row, axis=1)
    picked=[]
    for k,g in all_df.groupby("__key", dropna=True):
        best = g.sort_values(["Score","Sigma"], ascending=[False,False]).iloc[0].copy()
        srcs = set(map(str,g["__src"].tolist()))
        flags = str(best.get("Flags","")).strip()
        if flags in ("nan","None"): flags = ""
        flag_set = set([f.strip() for f in flags.replace("[","").replace("]","").replace("'","").split(",") if f.strip()])
        flag_set.add("src=both" if len(srcs)>1 else f"src={next(iter(srcs))}")
        best["Flags"] = ", ".join(sorted(flag_set)) if flag_set else ""
        picked.append(best)
    union = pd.DataFrame(picked)
    if "__src" in union.columns:
        def _with_src(row):
            t = str(row.get("Type","em"))
            flags = str(row.get("Flags",""))
            src = "both" if "src=both" in flags else ("band" if "src=band" in flags else ("xlsx" if "src=xlsx" in flags else None))
            if src and "·" not in t:
                return f"{t}·{src}"
            return t
        union["Type"] = union.apply(_with_src, axis=1)
    if "Score" in union.columns and "Sigma" in union.columns:
        union = union.sort_values(["Score","Sigma"], ascending=[False,False])
    return union

def _safe_copy(src, dst):
    try:
        if os.path.exists(src):
            shutil.copyfile(src, dst)
    except Exception:
        pass

def _write_docx_fallback(outdir, union_csv, band_img, graph_img, xlsx_img=None):
    try:
        import docx
        from docx.shared import Inches, RGBColor, Pt
    except Exception as e:
        with open(os.path.join(outdir,"_docx_error.txt"),"w",encoding="utf-8") as f:
            f.write(str(e))
        return

    doc = docx.Document()

    # superscript/subscript helper (måste definieras före användning)
    def _super_sub(txt: str) -> str:
        if not isinstance(txt, str): return txt
        for d, sub in zip("0123456789", "₀₁₂₃₄₅₆₇₈₉"):
            for el in ("H","He","C","N","O","Na","Mg","Al","Si","P","S","Cl","K","Ca","Fe","Li"):
                txt = txt.replace(el + d, el + sub)
        txt = (txt.replace("^-1","⁻¹").replace("^-2","⁻²").replace("^-3","⁻³")
                   .replace("^2","²").replace("^3","³"))
        return txt

    # Analys-sammanställning
    try:
        import pandas as _pd, numpy as _np
        doc.add_paragraph("Analys-sammanställning").runs[0].bold = True
        if union_csv and os.path.exists(union_csv):
            _df = _pd.read_csv(union_csv)
            dcol = next((c for c in _df.columns if c.lower() in ["delta_nm","Δλ","delta"]), None)
            med = None
            if dcol is not None:
                try: med = float(_np.median(_pd.to_numeric(_df[dcol], errors="coerce").dropna()))
                except Exception: med = None
            band_n = int((_df.get("Källa","")=="Band").sum()) if "Källa" in _df.columns else None
            data_n = int((_df.get("Källa","")=="Data").sum()) if "Källa" in _df.columns else None
            line = f"Totalt {len(_df)} linjer." + (f" Robust offset (median Δλ): {med:+.3f} nm." if med is not None else "")
            if band_n is not None and data_n is not None:
                line += f" Källa – Band: {band_n}, Data: {data_n}."
            doc.add_paragraph(_super_sub(line))
        doc.add_paragraph()
        doc.add_page_break()
    except Exception:
        pass

    # Instrument och kalibrering
    doc.add_paragraph("Instrument och kalibrering").runs[0].bold = True
    doc.add_paragraph(_super_sub(
        "Instrument: KVANT SPECTRA 1 – undervisningsspektrometer för VIS (~360–930 nm). "
        "Upplösning <1.8 nm FWHM; pixelupplösning <0.5 nm; gitter 500 linjer/mm; "
        "detektor 1/3″ färg-CMOS 1.3 MP (1280×960); Windows 8/10/11. Källa: forschool.eu / KVANT."
    ))
    doc.add_paragraph("Fabrikskalibrering (px→nm):")
    tcal = doc.add_table(rows=1, cols=2)
    tcal.rows[0].cells[0].text = "Pixel (px)"; tcal.rows[0].cells[1].text = "Våglängd (nm)"
    for _px, _nm in [(32,388.86),(515,587.57),(1110,837.76)]:
        _r = tcal.add_row().cells; _r[0].text=str(_px); _r[1].text=f"{_nm:.2f}"
    doc.add_paragraph()
    doc.add_page_break()

    # Compact layout default
    try:
        style = doc.styles['Normal']
        pf = style.paragraph_format
        pf.space_before = Pt(0); pf.space_after = Pt(0); pf.line_spacing = 1.0
    except Exception:
        pass

    doc.add_heading("Spektrumanalys – engine6 (union)", level=1)
    if band_img and os.path.exists(band_img):
        doc.add_paragraph("Band-first: annotated_band.png")
        try: doc.add_picture(band_img, width=Inches(6.5))
        except Exception: pass
    if graph_img and os.path.exists(graph_img):
        doc.add_paragraph("Band-first: annotated_graph.png")
        doc.add_paragraph()
        doc.add_page_break()
        try: doc.add_picture(graph_img, width=Inches(6.5))
        except Exception: pass
    if xlsx_img and os.path.exists(xlsx_img):
        doc.add_paragraph("XLSX-first: annotated_band.png")
        try: doc.add_picture(xlsx_img, width=Inches(6.5))
        except Exception: pass

    if union_csv and os.path.exists(union_csv):
        df = pd.read_csv(union_csv)
        if 'Type' in df.columns and 'Källa' not in df.columns:
            def _src_from_type(t):
                t = str(t)
                if '·both' in t: return 'Båda'
                if '·band' in t: return 'Band'
                if '·xlsx' in t: return 'Data'
                return ''
            df['Källa'] = df['Type'].apply(_src_from_type)
        table = doc.add_table(rows=1, cols=len(df.columns))
        hdr = table.rows[0].cells
        try:
            from docx.shared import RGBColor
            for j,c in enumerate(df.columns):
                run = hdr[j].paragraphs[0].add_run(str(c)); run.font.color.rgb=RGBColor(255,255,255)
        except Exception:
            for j,c in enumerate(df.columns):
                hdr[j].paragraphs[0].add_run(str(c))
        for _,row in df.iterrows():
            cells = table.add_row().cells
            for j,val in enumerate(row):
                cells[j].paragraphs[0].add_run(_super_sub(str(val)))

    try:
        from docx.shared import Pt
        for p in doc.paragraphs:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            p.paragraph_format.line_spacing = 1
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after  = Pt(0)
                        p.paragraph_format.line_spacing = 1
    except Exception:
        pass

    # v6.2.1 FIX: spara som report_fallback.docx (matchar rename-logiken)
    doc.save(os.path.join(outdir,"report_fallback.docx"))

def _make_band_original_adjusted(src_path, dst_path):
    try:
        from PIL import ImageEnhance
        img = Image.open(src_path).convert("RGB")
        img = ImageEnhance.Brightness(img).enhance(1.4)
        img = ImageEnhance.Contrast(img).enhance(0.6)
        img.save(dst_path)
        return dst_path
    except Exception:
        return None

def _find_investigation_image(band_path, graph_path):
    import glob
    base = os.path.dirname(band_path) if band_path else "."
    pats = ["IMG_*.*","img_*.*","DSC*.*","photo*.*","lamp*.*","tube*.*","unders*.*","setup*.*","*.jpeg","*.jpg"]
    cands = []
    for pat in pats:
        cands += glob.glob(os.path.join(base, pat))
    bad = set()
    for p in (band_path, graph_path):
        if p and os.path.exists(p):
            try: bad.add(os.path.realpath(p))
            except Exception: pass
    cands = [p for p in cands if (os.path.realpath(p) if os.path.exists(p) else p) not in bad]
    if not cands: return None
    cands.sort(key=lambda p: (-os.path.getsize(p), p.lower()) if os.path.exists(p) else (0,p.lower()))
    return cands[0]

def _union_csv_to_matches(union_csv_path):
    df = pd.read_csv(union_csv_path) if os.path.exists(union_csv_path) else pd.DataFrame()
    if df.empty: return []
    cols = {c.lower().strip(): c for c in df.columns}
    def col(*names):
        for n in names:
            if n in cols: return cols[n]
        return None
    c_species = col("element/species","species","element")
    c_ref     = col("ref_nm","refnm","reference_nm","ref")
    c_meas    = col("measured_nm","nm_meas","meas_nm","measured")
    c_corr    = col("corrected_nm","corr_nm","corrected")
    c_delta   = col("delta_nm","Δnm","delta")
    c_score   = col("score")
    c_sigma   = col("sigma")
    c_type    = col("type")
    c_rel     = col("relative_strength","rel","strength")
    c_flags   = col("flags")
    out = []
    for _, row in df.iterrows():
        m = {}
        if c_species: m["species"] = str(row[c_species])
        if c_ref and not pd.isna(row[c_ref]): m["ref_nm"] = float(row[c_ref])
        if c_meas and not pd.isna(row[c_meas]): m["nm_meas"] = float(row[c_meas])
        if c_corr and not pd.isna(row[c_corr]): m["corrected_nm"] = float(row[c_corr])
        if c_delta and not pd.isna(row[c_delta]):
            m["delta_nm"] = float(row[c_delta])
        else:
            try:
                m["delta_nm"] = float(m.get("corrected_nm", m.get("nm_meas", 0.0)) - m.get("ref_nm", 0.0))
            except Exception:
                pass
        if c_score and not pd.isna(row[c_score]): m["score"] = float(row[c_score])
        if c_sigma and not pd.isna(row[c_sigma]): m["sigma"] = float(row[c_sigma])
        base_t = "em" if not (c_type and not pd.isna(row[c_type])) else str(row[c_type])
        src_tag = None
        if c_flags and not pd.isna(row[c_flags]):
            txt = str(row[c_flags])
            if "src=both" in txt: src_tag = "both"
            elif "src=band" in txt: src_tag = "band"
            elif "src=xlsx" in txt: src_tag = "xlsx"
        m["type"] = f"{base_t}·{src_tag}" if src_tag and "·" not in base_t else base_t
        if c_rel and not pd.isna(row[c_rel]): m["rel_strength"] = float(row[c_rel])
        out.append(m)
    return out

def run_engine(dataset, band, graph_img_unused=None, mode="auto", k=7, threshold=0.04,
               min_dist_nm=1.2, tolerance_nm=3.0, outdir="./out", library_path=None):
    os.makedirs(outdir, exist_ok=True)
    out_xlsx = os.path.join(outdir,"_xlsx"); out_band=os.path.join(outdir,"_band")
    os.makedirs(out_xlsx, exist_ok=True); os.makedirs(out_band, exist_ok=True)

    df_xlsx=None
    if dataset and os.path.exists(dataset):
        base_run_engine(dataset=dataset, band=band, graph_img_unused=graph_img_unused, mode=mode, library_path=library_path,
                        k=k, threshold=threshold, min_dist_nm=min_dist_nm, tolerance_nm=tolerance_nm,
                        outdir=out_xlsx)
        p_x = os.path.join(out_xlsx,"lines.csv")
        if os.path.exists(p_x): df_xlsx = pd.read_csv(p_x)

    band_ds = os.path.join(outdir,"_dataset_from_band.csv")
    _build_dataset_from_band(band, band_ds)
    base_run_engine(dataset=band_ds, band=band, graph_img_unused=graph_img_unused, mode=mode, library_path=library_path,
                    k=k, threshold=threshold, min_dist_nm=min_dist_nm, tolerance_nm=tolerance_nm,
                    outdir=out_band)
    p_b = os.path.join(out_band,"lines.csv")
    df_band = pd.read_csv(p_b) if os.path.exists(p_b) else pd.DataFrame()

    if df_xlsx is not None:
        df_xlsx.to_csv(os.path.join(outdir,"lines_xlsx.csv"), index=False)
    if df_band is not None and not df_band.empty:
        df_band.to_csv(os.path.join(outdir,"lines_band.csv"), index=False)

    union = _union_lines(df_xlsx, df_band)
    union_csv = os.path.join(outdir,"lines.csv")
    union.to_csv(union_csv, index=False)

    _safe_copy(os.path.join(out_band,"annotated_band.png"), os.path.join(outdir,"annotated_band.png"))
    _safe_copy(os.path.join(out_band,"annotated_graph.png"), os.path.join(outdir,"annotated_graph.png"))
    _safe_copy(os.path.join(out_xlsx,"annotated_band.png"), os.path.join(outdir,"annotated_band_xlsx.png"))
    _safe_copy(os.path.join(out_xlsx,"annotated_graph.png"), os.path.join(outdir,"annotated_graph_xlsx.png"))

    try:
        orig_adj = os.path.join(outdir, 'band_original.png')
        _make_band_original_adjusted(band, orig_adj)
    except Exception:
        orig_adj = None

    try:
        INV = _find_investigation_image(band, graph_img_unused)
    except Exception:
        INV = None

    params = dict(mode=mode, k=k, threshold=threshold, min_dist_nm=min_dist_nm, tolerance_nm=tolerance_nm)
    qc_log = []
    try:
        lib_atoms = load_library(library_path) if library_path else []
    except Exception:
        lib_atoms = []

    meta = {"engine6": True,
            "params": {"mode": mode, "k": k, "threshold": threshold, "min_dist_nm": min_dist_nm, "tolerance_nm": tolerance_nm},
            "sources": {"dataset": dataset, "band": band, "graph": graph_img_unused},
            "counts": {"union": int(len(union)), "xlsx": 0 if df_xlsx is None else int(len(df_xlsx)),
                       "band": 0 if df_band is None else int(len(df_band))}}

    matches_for_doc = _union_csv_to_matches(union_csv)
    try:
        build_docx_v615(outdir,
                   os.path.join(outdir,"annotated_band.png"),
                   os.path.join(outdir,"annotated_graph.png"),
                   params, matches_for_doc, len(matches_for_doc),
                   meta, qc_log, lib_atoms, investigation_img=INV, original_band_path=orig_adj)
        try:
            import datetime as _dt
            _src = os.path.join(outdir, "report.docx")
            if os.path.exists(_src):
                _ts = _dt.datetime.now().strftime("%y%m%d%H%M%S")
                _dst = os.path.join(outdir, f"Report_{_ts}.docx")
                try: shutil.move(_src, _dst)
                except Exception: shutil.copyfile(_src, _dst)
        except Exception:
            pass
    except Exception as e:
        open(os.path.join(outdir,'_docx_exception.txt'),'w',encoding='utf-8').write(str(e))
        _write_docx_fallback(outdir, union_csv,
                             os.path.join(outdir,"annotated_band.png"),
                             os.path.join(outdir,"annotated_graph.png"),
                             xlsx_img=os.path.join(outdir,"annotated_band_xlsx.png"))
        try:
            import datetime as _dt
            _src = os.path.join(outdir, "report_fallback.docx")
            if os.path.exists(_src):
                _ts = _dt.datetime.now().strftime("%y%m%d%H%M%S")
                _dst = os.path.join(outdir, f"Report_{_ts}.docx")
                try: shutil.move(_src, _dst)
                except Exception: shutil.copyfile(_src, _dst)
        except Exception:
            pass

    with open(os.path.join(outdir,"_engine6_meta.json"),"w",encoding="utf-8") as f:
        f.write(json.dumps(meta, indent=2))
    return union_csv

# ------------------------------------------------------------------
#  Autorun-harness (”bara kör engine.py” i Notebook / KVANT-miljö)
# ------------------------------------------------------------------

if __name__ == "__main__" and not _ALLOW_CLI():  # v6.2.1 FIX: kör bara om inga --flaggar
    import csv as _csv
    from glob import glob
    from datetime import datetime as _dt

    DATA = "/mnt/data"
    OUT  = os.path.join(DATA, "out")
    os.makedirs(OUT, exist_ok=True)

    def pick_first(patterns, prefer):
        c=[]
        for pat in patterns: c.extend(glob(os.path.join(DATA, pat)))
        if not c: return None
        def score(path):
            base=os.path.basename(path).lower()
            for i,(sub,pr) in enumerate(prefer):
                if sub and sub in base: return pr*10+i
            return 9999
        c.sort(key=lambda p:(score(p), p.lower()))
        return c[0]

    ENGINE_PATH = __file__
    cands=[]
    for pat in ["image.*","spectrum.*","band.*","*.png","*.jpg","*.jpeg"]:
        cands += glob(os.path.join(DATA, pat))
    cands = [p for p in cands if not os.path.basename(p).lower().startswith("graph")]
    BAND=None
    if cands:
        pr=[("image",0),("spectrum",1),("band",2)]
        def bscore(p):
            base=os.path.basename(p).lower()
            for i,(tag,prv) in enumerate(pr):
                if tag in base: return prv*10+i
            return 50
        cands.sort(key=lambda p:(bscore(p), p.lower()))
        BAND=cands[0]

    GRAPH = pick_first(["Graph.*","graph.*","*.png","*.jpg","*.jpeg"], [("graph",0),("spectrum",9)])
    if GRAPH and BAND and os.path.samefile(GRAPH,BAND):
        GRAPH=None
    XDATA = pick_first(["*.xlsx","*.xls","*.csv"], [("values",0),("data",1),("",5)])

    assert BAND, "Hittar ingen spektrumbild (image_*, spectrum_* eller band_*)."

    print("Kör engine.py med:")
    print(" BAND :", os.path.basename(BAND))
    print(" GRAPH:", os.path.basename(GRAPH) if GRAPH else None)
    print(" XDATA:", os.path.basename(XDATA) if XDATA else None)

    run_engine(
        dataset=XDATA,
        band=BAND,
        graph_img_unused=GRAPH,
        mode="auto",
        k=7,
        threshold=0.04,
        min_dist_nm=1.2,
        tolerance_nm=3.0,
        outdir=OUT,
        library_path=os.path.join(DATA, 'line_library_v1.json') if os.path.exists(os.path.join(DATA, 'line_library_v1.json')) else '/mnt/data/line_library_v1.json',
    )

    ts = _dt.now().strftime("%y%m%d%H%M%S")
    for cand in ["report.docx","report_fallback.docx"]:
        p=os.path.join(OUT, cand)
        if os.path.exists(p):
            dst = os.path.join(OUT, f"report_{ts}.docx")
            try: os.replace(p, dst)
            except FileNotFoundError: pass
            else: print("Rapport:", dst)

    def read_csv_safe(path):
        with open(path,"r",newline="",encoding="utf-8") as f:
            sample=f.read(2048); f.seek(0)
            dialect=_csv.Sniffer().sniff(sample, delimiters=",;")
            rows=list(_csv.reader(f,dialect))
        return rows[0], rows[1:]

    merged_rows=[]; merged_header=None
    for path,src in [(os.path.join(OUT,"lines_band.csv"),"IMAGE"),
                     (os.path.join(OUT,"lines_xlsx.csv"),"XLSX")]:
        if os.path.exists(path):
            header,rows = read_csv_safe(path)
            if "source" not in [h.lower() for h in header]:
                header = header + ["source"]; rows=[r+[src] for r in rows]
            if merged_header is None: merged_header=header
            merged_rows.extend(rows)

    if merged_rows:
        out_csv=os.path.join(OUT,"lines_all.csv")
        with open(out_csv,"w",newline="",encoding="utf-8") as f:
            w=_csv.writer(f); w.writerow(merged_header)
            hl=[h.lower() for h in merged_header]
            sort_idx=None
            for key in ("wavelength_nm","lambda_nm","nm","wavelength","corrected_nm","ref_nm"):
                if key in hl: sort_idx=hl.index(key); break
            if sort_idx is not None:
                def to_float(s):
                    try: return float(str(s).replace(",",".").strip())
                    except: return 1e9
                merged_rows.sort(key=lambda r: to_float(r[sort_idx]))
            w.writerows(merged_rows)
        print("Sammanfogad linjetabell:", out_csv)

    for f in ["annotated_band.png","annotated_graph.png",
              "lines.csv","lines_band.csv","lines_xlsx.csv","lines_all.csv"]:
        p=os.path.join(OUT,f)
        if os.path.exists(p): print(" ->", p)
